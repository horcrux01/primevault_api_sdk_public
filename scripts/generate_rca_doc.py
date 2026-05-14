from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color,
        qn("w:val"): "clear",
    })
    shading.append(shading_elem)


# ============================================================
# Title
# ============================================================
title = doc.add_heading(
    "RCA \u2014 API Signing Key Exposure via Public Branch Push", level=0
)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

# ============================================================
# Metadata table
# ============================================================
meta_table = doc.add_table(rows=3, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ("Date of Incident", "April 10, 2026"),
    ("Severity", "Critical"),
    ("Status", "Contained \u2014 remediation in progress"),
]
for i, (label, value) in enumerate(meta_data):
    meta_table.rows[i].cells[0].text = label
    meta_table.rows[i].cells[1].text = value
    for cell in meta_table.rows[i].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)
    meta_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    set_cell_shading(meta_table.rows[i].cells[0], "E8EAF6")

doc.add_paragraph()

# ============================================================
# 1. Summary
# ============================================================
add_heading("1. Summary")
add_para(
    "A test API signing key was inadvertently exposed on a public GitHub branch "
    "when an AI coding agent pushed local branch code that included the key in "
    "test case files. The branch itself was publicly accessible. An external "
    "attacker discovered the exposed key and attempted unauthorized transactions."
)

# ============================================================
# 2. Timeline
# ============================================================
add_heading("2. Timeline of Events")
timeline = [
    ("6th April, 2026",
     "AI agent pushes a branch to primevault_api_sdk_public containing API key in test case files"),
    ("10th April, 2026",
     "Attacker discovers the exposed API key on the public branch"),
    ("10th April, 2026 ~8:10 p.m.",
     "Attempt 1: Attacker initiates a $1 transfer \u2014 succeeds (no approval required for small amount)"),
    ("10th April, 2026 ~8:17 p.m.",
     "Attempt 2: Attacker initiates a $1,000 transfer \u2014 fails (requires policy approval, nobody approves)"),
    ("10th April, 2026 ~8:20 p.m.",
     "Attempt 3: Attacker initiates a $45,000 transfer \u2014 blocked"),
    ("10th April, 2026 ~8:20 p.m.",
     'Vivek notices unusual activity and asks Tushal: "Are you using your API user for this?"'),
    ("10th April, 2026 ~8:20 p.m.",
     "Tushal confirms he is not \u2014 team recognises the account is under attack"),
    ("10th April, 2026 ~8:30 p.m.",
     "All access tied to Tushal\u2019s user/email is immediately revoked"),
]
table = doc.add_table(rows=len(timeline) + 1, cols=2)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0]
hdr.cells[0].text = "Time"
hdr.cells[1].text = "Event"
for cell in hdr.cells:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)
    set_cell_shading(cell, "1A1A2E")
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for i, (time_str, event) in enumerate(timeline):
    row = table.rows[i + 1]
    row.cells[0].text = time_str
    row.cells[1].text = event
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)

doc.add_paragraph()

# ============================================================
# 3. Impact
# ============================================================
add_heading("3. Impact")
add_bullet("~$1 (single successful micro-transfer before detection)", bold_prefix="Financial loss: ")
add_bullet("$46,000+ (blocked by approval policy and quick response)", bold_prefix="Prevented loss: ")
add_bullet("API signing key for one user \u2014 revoked immediately", bold_prefix="Credential scope: ")
add_bullet("No customer data was accessed or exfiltrated", bold_prefix="Data exposure: ")

# ============================================================
# 4. Root Cause (5 reasons)
# ============================================================
add_heading("4. Root Cause")

add_bullet(
    "The repository had no pre-commit hooks or CI pipeline checks to detect "
    "secrets before they reached a public branch.",
    bold_prefix="1. No secret-scanning in CI or pre-commit \u2014 ",
)
add_bullet(
    "The AI coding agent had direct write access to the public repository and "
    "pushed a branch containing test credentials without human review.",
    bold_prefix="2. AI agent pushed directly to a public repo \u2014 ",
)
add_bullet(
    "Even though the PR diff was clean, the underlying branch (with test files "
    "containing the key) was publicly browsable on GitHub. Any visitor could "
    "view the branch contents.",
    bold_prefix="3. Branch visibility on public repos \u2014 ",
)
add_bullet(
    "There was no separation between development (private) and release (public) "
    "repositories. Code with test credentials was pushed to the same repo that "
    "external users can access.",
    bold_prefix="4. No private/public repo separation \u2014 ",
)
add_bullet(
    "The API key used in test cases was a real, active signing key rather than a "
    "dummy or environment-scoped token. A test-only key with no transaction "
    "permissions would have rendered the exposure harmless.",
    bold_prefix="5. Real API key used in test fixtures \u2014 ",
)

# ============================================================
# 5. Why the Approval Policy Saved Us
# ============================================================
add_heading("5. Why the Approval Policy Saved Us")
add_para(
    "PrimeVault\u2019s transaction approval policy requires multi-party sign-off "
    "for transfers above a configurable threshold. The attacker could only execute "
    "a $1 transfer autonomously. Both the $1,000 and $45,000 attempts required "
    "approvals that were never granted, buying the team time to detect and respond."
)

# ============================================================
# 6. Immediate Actions Taken
# ============================================================
add_heading("6. Immediate Actions Taken")

add_heading("6.1  Credential Revocation", level=2)
add_bullet("All API keys, sessions, and access tied to the compromised user were revoked immediately.")
add_bullet("Verified no other credentials were exposed in the repository history.")

add_heading("6.2  Gitleaks \u2014 Pre-commit & CI Secret Scanning", level=2)
add_bullet("Added gitleaks as a pre-commit hook across all SDK repos (Python & JS).")
add_bullet("Added gitleaks as a GitHub Actions workflow that runs on every push and PR.")
add_bullet("Known false positives (test UUIDs, example keys) are tracked in .gitleaksignore files.")

add_heading("6.3  Private Repo Workflow", level=2)
add_bullet("Created private mirror repos for both Python and JS SDKs.")
add_bullet("All development now happens in private repos.")
add_bullet(
    "A manual-trigger GitHub Action (sync-to-public) pushes changes from private "
    "main to a PR on the public repo \u2014 no direct pushes to public."
)
add_bullet("Workflow files from the private repo are excluded from sync to prevent leakage.")

add_heading("6.4  NPM Version Monitoring", level=2)
add_bullet(
    "Added a scheduled GitHub Action that runs hourly to compare package.json "
    "version against the npm registry."
)
add_bullet("Any mismatch triggers a Slack alert \u2014 catches unauthorized publishes.")

add_heading("6.5  IAM Role Separation", level=2)
add_bullet("Created separate IAM roles with least-privilege access.")
add_bullet("Python backend roles have no access to TSS or KMS.")
add_bullet(
    "Agent (AI) access is restricted \u2014 limited to pushing branches/PRs, "
    "no direct main branch access."
)

# ============================================================
# 7. Next Steps
# ============================================================
add_heading("7. Next Steps")

add_heading("7.1  Session Activity Logging", level=2)
add_bullet("Implement detailed session activity logging for all API users.")
add_bullet("Log every authentication event, transaction initiation, and approval action.")

add_heading("7.2  Slack Alerts on Privileged IAM Usage", level=2)
add_bullet("Set up real-time Slack notifications when IAM roles with elevated access are used.")
add_bullet("Covers: TSS access, KMS access, admin operations.")

add_heading("7.3  Deprecate Broad IAM Access", level=2)
add_bullet("Migrate all write operations to a super-user page with explicit audit trail.")
add_bullet(
    "Deprecate current IAM roles that have write access once the super-user "
    "migration is complete."
)

add_heading("7.4  Move Write Operations to Enclave", level=2)
add_bullet("Write operations will be deprecated from the current architecture.")
add_bullet("All sensitive write operations will move to an enclave-based execution model.")

add_heading("7.5  Alerts on User & Policy Changes", level=2)
add_bullet(
    "Any addition of a new user or modification of an approval policy must "
    "trigger an alert."
)
add_bullet("Alert channels: Slack + email to security team.")

# ============================================================
# 8. Lessons Learned
# ============================================================
add_heading("8. Lessons Learned")
add_bullet(
    "Branch content is as visible as merged code on public repositories.",
    bold_prefix="Never push secrets to public repos, even on branches \u2014 ",
)
add_bullet(
    "Agents with repo write access must be scoped to private repos only, with "
    "human review before anything reaches public.",
    bold_prefix="AI agents need guardrails \u2014 ",
)
add_bullet(
    "The multi-party approval requirement prevented a $46K loss and gave the "
    "team time to respond.",
    bold_prefix="Approval policies are a critical safety net \u2014 ",
)
add_bullet(
    "No single control stopped the attack; it was the combination of transaction "
    "limits, approval policies, and human vigilance.",
    bold_prefix="Defense in depth works \u2014 ",
)
add_bullet(
    "Manual review cannot reliably catch every secret; pre-commit hooks and CI "
    "scanning are essential.",
    bold_prefix="Automate secret detection \u2014 ",
)

# ============================================================
# Save
# ============================================================
out_path = os.path.expanduser("~/primevault_api_sdk_public/docs/RCA-API-KEY-EXPOSURE.docx")
doc.save(out_path)
print(f"Saved to {out_path}")
